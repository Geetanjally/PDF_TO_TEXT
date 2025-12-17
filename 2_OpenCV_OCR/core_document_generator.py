import json
import io
import time
import os
import base64
from google import genai
from google.genai import Client
from google.genai.types import Content, Part, GenerateContentConfig
from PIL import Image

# --- Conditional Imports for Libraries used in the advanced pipeline ---
try:
    from pptx_designer import create_pptx_with_style as create_pptx
except ImportError:
    create_pptx = None
    
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import fitz # PyMuPDF for PDF text extraction and image rendering
except ImportError:
    fitz = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    import cv2
    import numpy as np
except ImportError:
    cv2 = None
    np = None

try:
    from google.genai import Client
    from google.genai.types import GenerateContentConfig
except ImportError:
    Client = None
    GenerateContentConfig = None

# --------------------------------
# CONFIGURATION
# --------------------------------
GENAI_CLIENT = None
# --------------------------------

# --------------------------------
# 2. CORE GEMINI API FUNCTIONS
# --------------------------------

def _get_genai_client(api_key):
    """Initializes and returns the global Gemini API client."""
    global GENAI_CLIENT
    if GENAI_CLIENT is None:
        if not api_key:
            raise ValueError("Gemini API Key is required.")
        GENAI_CLIENT = Client(api_key=api_key)
    return GENAI_CLIENT

def _retry_api_call(func, *args, **kwargs):
    """Implements exponential backoff for API calls."""
    max_retries = 5
    for attempt in range(max_retries):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            if attempt < max_retries - 1:
                delay = 2 ** attempt
                print(f"API Error: {e}. Retrying in {delay} seconds...")
                time.sleep(delay)
            else:
                raise

def extract_text_gemini(image_bytes, api_key):
    """
    Uses the Gemini API with vision capabilities to extract text from an image.
    
    Args:
        image_bytes (bytes): Byte content of the image (e.g., PNG).
        api_key (str): The Gemini API key.
        
    Returns:
        str: The extracted text, or None if extraction fails.
    """
    try:
        client = _get_genai_client(api_key)
    except ValueError as e:
        return None, str(e)
        
    parts = [
        # 1. Image part
        Part.from_bytes(data=image_bytes, mime_type="image/png"),
        
        # 2. Text instruction part
        Part.from_text(text="Extract ALL text accurately from this image. Preserve line breaks and formatting. Do NOT summarize or add commentary. Return ONLY the raw text."),
    ]
    
    try:
        response = _retry_api_call(
            client.models.generate_content,
            model="gemini-flash-latest",
            contents=parts,
            config=GenerateContentConfig(temperature=0.0)
        )
        return response.text, None
        
    except Exception as e:
        return None, f"Gemini API call failed during text extraction: {e}"


def process_document_to_cleaned_text(uploaded_file, api_key):
    """
    Processes an uploaded file (PDF/Image) to extract raw text, 
    then uses Gemini to clean and structure it.
    """
    raw_text = None
    
    if uploaded_file.type == "application/pdf":
        pdf_bytes = uploaded_file.getvalue()
        
        if fitz:
            try:
                doc = fitz.open(stream=pdf_bytes)
                
                # **********************************************
                # CASE 1: TEST ONLY - Single Page Processing
                # **********************************************
                # Uncomment the block below (remove the triple quotes) 
                # to process ONLY the first page for fast, low-token testing.
                # **********************************************
                '''
                page_num = 0 # Only process the first page (index 0)
                page = doc.load_page(page_num)
                
                # Convert page to a high-resolution PNG image
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")
                
                # Send image to Gemini Vision for comprehensive text extraction
                page_text, error = extract_text_gemini(img_bytes, api_key)
                
                if error:
                    page_text = f"[ERROR: Could not extract page 1]. {error}"
                    
                raw_text = f"\n\n--- PAGE 1 ---\n\n{page_text}"
                print(f"📄 Testing mode: Only Page 1 extracted.")
                '''
                # Ensure you comment out the full scanning block (CASE 2) if using this!
                
                
                # **********************************************
                # CASE 2: FINAL SUBMISSION - Whole PDF with Optimization
                # **********************************************
                # Use this block for the final submission. It includes the token-saving logic.
                # **********************************************
                
                raw_text_parts = []
                for page_num in range(doc.page_count):
                    page = doc.load_page(page_num)
                    page_text = ""
                    
                    # 1. Try to extract digital text first (LOW COST / FAST)
                    raw_pdf_text = page.get_text()
                    meaningful_text_length = len(raw_pdf_text.strip())
                    
                    # Optimization: If digital text is substantial, use it to save API cost.
                    if meaningful_text_length > 250:
                        page_text = raw_pdf_text.strip()
                        print(f"📄 Page {page_num + 1}: Digital text found, skipping Gemini OCR.")
                    
                    # 2. Fallback: If digital text is sparse or missing, use Gemini Vision (HIGH COST)
                    else:
                        print(f"🖼️ Page {page_num + 1}: Digital text sparse ({meaningful_text_length} chars). Falling back to Gemini Vision OCR...")
                        
                        # Convert page to high-resolution PNG image
                        pix = page.get_pixmap(dpi=300)
                        img_bytes = pix.tobytes("png")
                        
                        # Send image to Gemini Vision for comprehensive text extraction
                        page_text, error = extract_text_gemini(img_bytes, api_key)
                        
                        if error:
                            print(f"Warning: Failed to process page {page_num+1}. Error: {error}")
                            page_text = f"[ERROR: Could not extract page {page_num+1}]" 
                        
                    raw_text_parts.append(f"\n\n--- PAGE {page_num + 1} ---\n\n{page_text}")
                    
                raw_text = "".join(raw_text_parts) # Combine all page text
            
                
                # Ensure you comment out the single-page block (CASE 1) if using this!
                
                doc.close()
                
                if not raw_text.strip():
                     return None, "PDF processed successfully but returned empty content."
                
            except Exception as e:
                return None, f"PDF Processing Error: {e}"
        else:
            return None, "PyMuPDF (fitz) library is missing, cannot process PDF."
            
    else: # Image file (standard image processing)
        img_bytes = uploaded_file.getvalue()
        raw_text, error = extract_text_gemini(img_bytes, api_key)
        if error:
            return None, f"Extraction Error (Image): {error}"

    # Check if raw text was successfully extracted before proceeding to cleaning
    if not raw_text:
        return None, "Raw text extraction failed or returned empty content."

    # Step 2: Use Gemini to clean and structure the extracted raw text
    clean_prompt = (
        "The following text was extracted from a document. "
        "Review the text and generate a structured, clean JSON list suitable for a presentation. "
        "The JSON MUST follow this exact schema: "
        "[{'title': 'Slide Title 1', 'content': ['Point 1', 'Point 2', '...', 'Point N']}, "
        "{'title': 'Slide Title 2', 'content': ['Point 1', 'Point 2', '...']}, ...]. "
        "Infer logical slide breaks from the content (e.g., section headings, major topic changes). "
        "Use simple bullet points in the 'content' lists. Do not use Markdown formatting in the lists."
        f"\n\nRAW TEXT:\n---\n{raw_text}\n---"
    )
    
    # ... (rest of the function for API calls remains the same)
    
    try:
        client = _get_genai_client(api_key)
    except ValueError as e:
        return None, str(e)
        
    # --- Structured JSON generation call ---
    try:
        response = _retry_api_call(
            client.models.generate_content,
            model="gemini-flash-latest",
            contents=clean_prompt,
            config=GenerateContentConfig(
                response_mime_type="application/json",
                response_schema={
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string", "description": "The title of the presentation slide."},
                            "content": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "A list of bullet points for the slide content."
                            }
                        },
                        "required": ["title", "content"]
                    }
                }
            )
        )
        return response.text, None
        
    except Exception as e:
        return None, f"Gemini API call failed during structuring/cleaning: {e}"
# --------------------------------
# 3. STRUCTURE MANIPULATION FUNCTIONS
# --------------------------------

def generate_initial_structure(input_text, system_instruction, api_key):
    """
    Generates an initial structured JSON for a given topic or existing text.
    """

    # If input_text is very short, assume it's a topic
    if len(input_text.split()) < 10:
        prompt_content = f"on the topic: '{input_text}'. "
    else:
        # Otherwise assume it is full cleaned text
        prompt_content = f"based on the following content:\n\n{input_text}\n\n"

    # Build the prompt
    prompt = (
        f"{system_instruction}\n\n"
        f"Create a structured JSON outline for a presentation {prompt_content}"
        "The JSON must follow this strictly:\n"
        "[\n"
        "  {\n"
        '    "title": "Slide Title",\n'
        '    "content": ["Point 1", "Point 2", "..."]\n'
        "  }\n"
        "]"
    )

    
    try:
        client = _get_genai_client(api_key)
    except ValueError as e:
        return None, str(e)

    try:
        response = _retry_api_call(
            client.models.generate_content,
            model="gemini-flash-latest",
            contents=prompt,
            config=GenerateContentConfig(
                response_mime_type="application/json",
                response_schema={
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string", "description": "The title of the presentation slide."},
                            "content": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "A list of bullet points for the slide content."
                            }
                        },
                        "required": ["title", "content"]
                    }
                }
            )
        )
        return response.text, None
    except Exception as e:
        return None, f"Gemini API call failed during structure generation: {e}"

def update_structure(api_key, existing_json, user_prompt):
    """Updates the existing JSON structure based on a user prompt."""
    
    prompt = (
        "You are an editor modifying a presentation structure. "
        "The following is the current structure (JSON) and a user's request. "
        "Modify the structure based on the request and return the new, complete JSON structure. "
        "Strictly adhere to the JSON schema: "
        "[{'title': 'Slide Title 1', 'content': ['Point 1', 'Point 2', '...']}, ...]. "
        f"\n\nCURRENT JSON:\n---\n{existing_json}\n---\n\nUSER REQUEST: {user_prompt}"
    )

    try:
        client = _get_genai_client(api_key)
    except ValueError as e:
        return None, str(e)
        
    try:
        response = _retry_api_call(
            client.models.generate_content,
            model="gemini-flash-latest",
            contents=prompt,
            config=GenerateContentConfig(
                response_mime_type="application/json",
                response_schema={
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string", "description": "The title of the presentation slide."},
                            "content": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "A list of bullet points for the slide content."
                            }
                        },
                        "required": ["title", "content"]
                    }
                }
            )
        )
        return response.text, None
    except Exception as e:
        return None, f"Gemini API call failed during structure update: {e}"


# --------------------------------
# 4. OUTPUT GENERATION FUNCTIONS (PPTX, DOCX, MD)
# --------------------------------

# NOTE: create_pptx is aliased to create_pptx_with_style from pptx_designer.py
if create_pptx is None:
    def create_pptx(slides_data, template_data=None):
        """Placeholder for PPTX creation if the required library is missing."""
        return None, "The 'python-pptx' library is not installed or the advanced designer could not be imported."


def create_docx(slides_data):
    """Generates a DOCX file from JSON structure."""
    if Document is None:
        return None, "The 'python-docx' library is not installed."
        
    try:
        data = json.loads(slides_data)
    except:
        return None, "Failed to parse blueprint data for DOCX. Check the JSON format."
        
    doc = Document()

    for i, entry in enumerate(data):
        title = entry.get('title', f"Slide {i+1}")
        content = entry.get('content', [])

        # Title/Heading
        if i == 0:
            doc.add_heading(title, level=1)
        else:
            doc.add_heading(title, level=2)

        # Content bullets
        for point in content:
            doc.add_paragraph(point, style='List Bullet') 

        if i < len(data) - 1:
            doc.add_page_break()

    # Save to a BytesIO buffer
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream, None

def create_markdown_report(slides_data):
    """Generates a detailed Markdown report from the JSON structure."""
    try:
        data = json.loads(slides_data)
    except:
        return None, "Failed to parse blueprint data for Markdown. Check the JSON format."

    markdown = "# Presentation Content Report\n\n"
    
    for i, entry in enumerate(data):
        title = entry.get('title', f"Slide {i+1}")
        content = entry.get('content', [])

        markdown += f"## {title}\n"
        if content:
            for point in content:
                # Use standard markdown bullet points
                markdown += f"- {point}\n"
        markdown += "\n"

    return markdown, None